import time 
import os
import re 
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from google import genai
import json

load_dotenv()

API_KEYS = [
    os.getenv("GEMINI_API_KEY_1"),
    os.getenv("GEMINI_API_KEY_2")
]
current_key = 0
clients = [
    genai.Client(api_key=key)
    for key in API_KEYS
]

def narrative_architect(topic):
    prompt = f"""
You are a literary consultant.

For this story concept:

{topic}

Select:

- Narrative perspective
- Tense
- Narrative distance
- Voice style
- Structural device

Choose what best serves the story.

Possible perspectives:
- First person
- Third person limited
- Third person omniscient
- Multiple POV
- Epistolary
- Documentary
- Unreliable narrator

Return ONLY valid JSON.

{
    "perspective":"",
    "tense":"",
    "distance":"",
    "voice":"",
    "structure":""
}
"""
    return generate(prompt)

def final_editor(story, narrative):

    constraints = ""

    for k, v in narrative.items():
        constraints += f"{k}: {v}\n"

    prompt = f"""
You are a professional developmental editor.

Narrative blueprint:

{constraints}

Your job:

- Preserve the narrative blueprint.
- Improve pacing.
- Improve continuity.
- Improve emotional arcs.
- Improve character consistency.
- Improve thematic depth.
- Do NOT alter major plot events.
- Do NOT change the chosen narrative perspective.

Special instruction:
If the narrative blueprint specifies a gradual emotional progression,
ensure that progression remains consistent across the entire novel.

Story:

{story}
"""

    return generate(prompt)

def txt_to_docx_kdp(input_path, output_path, title, author):
    document = Document()

    section = document.sections[0]

    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    section.left_margin = Inches(0.9)   
    section.right_margin = Inches(0.7)  

    style = document.styles["Normal"]

    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # ---- Title Page ----
    title_paragraph = document.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(28)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()  # spacer

    author_paragraph = document.add_paragraph()
    author_run = author_paragraph.add_run(f"By {author}")
    author_run.font.name = "Times New Roman"
    author_run.font.size = Pt(16)
    author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()

    document.add_page_break()

    copyright_page = document.add_paragraph()
    copyright_page.alignment = WD_ALIGN_PARAGRAPH.CENTER

    copyright_page.add_run(
        f"Copyright © 2026 {author}\n\n"
        "All rights reserved.\n\n"
        "This is a work of fiction.\n"
        "Names, characters, places,\n"
        "and events are products\n"
        "of the author's imagination."
    )

    document.add_page_break()

    # ---- Body ----
    with open(input_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    chapter_count = 1
    for line in lines:
        stripped = line.strip()

        if not stripped:
            continue

        # Optional: treat blank lines between sections as chapters
        if stripped.lower().startswith("chapter"):
            document.add_page_break()
            for _ in range(8):
                document.add_paragraph()
            heading = document.add_heading(level=1)
            run = heading.add_run(stripped)
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(18)

            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped == "***":
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("***")
            run.font.size = Pt(14)
            run.bold = True
            continue

        else:
            paragraph = document.add_paragraph(stripped)
            fmt = paragraph.paragraph_format
            fmt.first_line_indent = Inches(0.3)
            fmt.space_after = Pt(0)
            fmt.line_spacing = 1.15
    # ---- Save File ----
    document.save(output_path)

def extract_score(feedback):
    """
    Extracts a score (1–10) from model feedback text.
    Returns integer score or 0 if not found.
    """

    # 1. Look for explicit "Score: X" style
    match = re.search(r"score\s*[:\-]?\s*(\d{1,2})\s*(?:/10)?",
                      feedback,
                      re.IGNORECASE)

    if match:
        score = int(match.group(1))
        return max(1, min(score, 10))

    # 2. Fallback: find standalone number 1–10
    match = re.search(r"\b(10|[1-9])\b", feedback)
    if match:
        return int(match.group(1))

    return 0

def generate(prompt,max_retries=5):
    global current_key
    for attempt in range(max_retries):  
        try:
            client = clients[current_key]
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=prompt
            )
            return response.text

        except Exception as e:
            print(e)
            if "RESOURCE_EXHAUSTED" in str(e):
                current_key += 1
                if current_key >= len(clients):
                    raise Exception(
                        "All API keys exhausted."
                    )
                print(
                    f"Switching to API key {current_key+1}"
                )
                
                sleep_time = 30 * (2 ** attempt)
                print(f"Rate limit hit. Sleeping for {sleep_time} seconds before retrying...")
                time.sleep(sleep_time)
            else:
                time.sleep(30)

    raise Exception("API failed ")

#For creating the basic architecture of the story 
def story_architect(topic, total_chapters=8):
    prompt = f"""
You are a professional story architect.

Topic:
{topic}

Create a complete story blueprint consisting of exactly
{total_chapters} chapters.

For each chapter provide:

- Chapter number
- Chapter title
- Goal
- Main conflict
- Emotional state of protagonist
- Important revelation
- Ending hook

Requirements:

- The story must escalate naturally.
- The climax should occur near the end.
- The final chapter must fully resolve the story.
- Maintain thematic consistency.

Return ONLY valid JSON:

{
  "chapters":[
    {
      "number":1,
      "title":"",
      "goal":"",
      "conflict":"",
      "emotion":"",
      "revelation":"",
      "hook":""
    }
  ]
}

"""

    return generate(prompt)

#Long term memory to remember the story state
def update_memory(previous_memory, new_chapter):

    prompt = f"""
You are the story memory manager.

Existing memory:

{previous_memory}

New chapter:

{new_chapter}

Update the memory database.

Return exactly:

SUMMARY:
A concise summary of the story so far.

CHARACTER_STATE:
Current emotional and psychological states.

OPEN_THREADS:
Unresolved mysteries.

IMPORTANT_OBJECTS:
Important items and locations.

UNRESOLVED_CONFLICTS:
Conflicts still active.

THEMATIC_PROGRESS:
How the theme is evolving.

Keep the memory under 500 words.
"""

    return generate(prompt)

def parse_story_memory(memory_text):

    try:
        return json.loads(memory_text)

    except:
        return {
            "summary":"",
            "character_states":{},
            "open_threads":[],
            "important_objects":[],
            "conflicts":[],
            "theme_progress":""
        }
    
def safe_json_parse(text):

    text = text.strip()

    text = text.replace("```json","")
    text = text.replace("```","")

    return json.loads(text)


def story_foundation(topic):

    prompt = f"""
You are designing a story foundation.

Topic:
{topic}

Generate:

WORLD:
Describe the setting.

CHARACTERS:
Create 2-4 major characters.

THEME:
Define the emotional and philosophical theme.

STYLE:
Define writing style, tone, narration, pacing.

TITLE:
Generate a Title that will be used for the story based on the topic and other things.

Keep everything concise but detailed.
Return ONLY valid JSON.

{{
    "title":"",
    "world":"",
    "characters":[
        {{
            "name":"",
            "role":"",
            "motivation":""
        }}
    ],
    "theme":"",
    "style":""
}}
"""

    return generate(prompt)

def consistency_agent(
        chapter,
        world,
        characters,
        memory,
        narrative):

    constraints = ""

    for k,v in narrative.items():
        constraints += f"{k}: {v}\n"

    prompt = f"""
You are a story continuity editor.

Narrative blueprint:
{constraints}

World:
{world}

Characters:
{characters}

Story memory:
{memory}

Current chapter:
{chapter}

Check for:

- character inconsistencies
- timeline problems
- worldbuilding contradictions
- broken emotional arcs
- forgotten plot threads
- narration inconsistencies
- style inconsistencies

Return:

CONSISTENT: YES/NO

ISSUES:
- ...

FIXES:
- ...
"""

    return generate(prompt)

def parse_consistency(report):

    if "CONSISTENT: YES" in report:
        return True

    return False


def chapter_writer(
        chapter_info,
        world,
        characters,
        theme,
        style,
        memory,
        narrative):
    
    prompt = f"""
World:
{world}

Characters:
{characters}

Theme:
{theme}

Style:
{style}

Story memory:
{memory}

Use the memory database as the authoritative
source of continuity information.
Do not invent events that contradict it.

Current chapter blueprint:
{chapter_info}

Narrative blueprint:
{narrative}

Write ONLY this chapter.

Requirements:

- Produce approximately 2500 words.
- Maintain continuity.
- Follow the chapter blueprint exactly.
- End using the specified hook.
- Do not summarize future chapters.
"""
    
    return generate(prompt)

def consistency_fixer(chapter, report):

    prompt = f"""
Fix this chapter using the consistency report.

Report:
{report}

Chapter:
{chapter}

Preserve plot and writing quality.
Only fix inconsistencies.
"""

    return generate(prompt)

def chapter_critic(chapter, narrative):
    prompt = f"""
Critique this section briefly.
Give Score: X (1–10)
 
Section:
{chapter}

Narrative requirements:
{narrative}

Evaluate:

- narrative consistency
- pacing
- emotional progression
- voice consistency
- characterization

Return:
Score: X/10
Strengths:
Weaknesses:
"""
    return generate(prompt)

def log(text):
    with open("stories/log.txt", "a") as f:
        f.write(text + "\n\n")

def build_story(topic, story_id):
    i = 0
    # 1. Build foundation
    print("building the chapter plan")
    chapter_plan = safe_json_parse(story_architect(topic,8))
    print("Developing the world, characters, theme and style")
    foundation = safe_json_parse(story_foundation(topic))
    world = foundation["world"]
    characters = foundation["characters"]
    theme = foundation["theme"]
    style = foundation["style"]
    title = foundation["title"]
    chapters = []
    print("Generating a narrative agent")
    narrative = safe_json_parse(narrative_architect(topic))
    story_memory = """
        SUMMARY:

        CHARACTER_STATE:

        OPEN_THREADS:

        IMPORTANT_OBJECTS:

        UNRESOLVED_CONFLICTS:

        THEMATIC_PROGRESS:
         """
    memory_context = ""

    for i, chapter in enumerate(chapter_plan["chapters"]):

        print(f"Writing chapter {i+1}")

        generated = chapter_writer(
            chapter_info=chapter,
            world=world,
            characters=characters,
            theme=theme,
            style=style,
            memory=memory_context,
            narrative=narrative
        )

        feedback = chapter_critic(generated, narrative)

        score = extract_score(feedback)

        if score < 5:
            generated = chapter_writer(
                chapter_info=chapter,
                world=world,
                characters=characters,
                theme=theme,
                style=style,
                memory=memory_context,
                narrative=narrative
            )

        # if (i + 1) in [3, 6, 8]:
        #     consistency_report = consistency_agent(
        #         generated,
        #         world,
        #         characters,
        #         story_memory,
        #         narrative
        #     )

        #     if not parse_consistency(consistency_report):
        #         generated = consistency_fixer(
        #             generated,
        #             consistency_report
        #         )

        memory_context = f"""
            Summary:
            {story_memory['summary']}

            Character States:
            {story_memory['character_states']}

            Open Threads:
            {story_memory['open_threads']}

            Objects:
            {story_memory['important_objects']}

            Conflicts:
            {story_memory['conflicts']}

            Theme:
            {story_memory['theme_progress']}
            """
        
        with open(
            f"stories/story_{title}_memory.json",
            "w"
        ) as f:
            json.dump(
                story_memory,
                f,
                indent=4
            )

    metadata = {
    "title": title,
    "topic": topic,
    "world": world,
    "characters": characters,
    "theme": theme,
    "style": style,
    "narrative": narrative,
    "outline": chapter_plan,
    "story_memory": memory_context,
    "chapter_count":len(chapters)
    }

    with open(f"stories/story_{title}_meta.json", "w") as f:
        json.dump(metadata, f, indent=4)    

    final_story = "\n\n".join(chapters)
    final_story =final_editor(final_story, narrative)

    with open(f"stories/story_{story_id}.txt", "w") as f:
        print("Inside the file")
        f.write(final_story)

    return title, final_story
        
def kdp_formatter_agent(story):

    prompt = f"""
Format this novel for Kindle Direct Publishing.

Rules:

- Insert CHAPTER headers.
- Insert scene breaks using ***
- Preserve story content.
- Do not rewrite prose.
- Return plain text.

Story:

{story}
"""
    return generate(prompt)

def main():
    topic = """In a near-future metropolis, the city infrastructure quietly edits citizens’ memories to maintain social harmony. Minor heartbreak? Deleted. Political anger? Softened. One archivist discovers her own childhood has been rewritten dozens of times and starts restoring forbidden memories across the population. Theme fuel: identity vs comfort
Hook: The antagonist isn’t evil. It’s municipal optimization."""
    story_count = 0

    while story_count<1:
        try:
            title, final_story = build_story(topic, story_count)
            formatted_story = kdp_formatter_agent(final_story)

            with open(f"stories/story_{title}.txt","w",encoding="utf-8") as f:
                f.write(formatted_story)
        except Exception as e:
            log(str(e))
            time.sleep(5)
        
        txt_path = f"stories/story_{title}.txt"

        txt_to_docx_kdp(
            input_path=txt_path,
            output_path=f"stories/story_{title}.docx",
            title=title,
            author="Darshith Shetty"
        )

        story_count += 1
        time.sleep(10)

    print("The story is ready ")
    
if __name__ == "__main__":
    main()
